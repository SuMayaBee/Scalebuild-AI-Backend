from services.storage_service import upload_to_gcs
import io
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from PIL import Image
import tempfile
from bs4 import BeautifulSoup
import re

async def save_document_to_gcs(document_content: str, document_type: str, generated_for: str):
    """Save document to Google Cloud Storage as a text file"""
    try:
        # Create filename
        safe_name = "".join(c for c in generated_for if c.isalnum() or c in (' ', '-', '_')).strip()
        file_name = f"{document_type.replace(' ', '_')}_{safe_name.replace(' ', '_')[:30]}.txt"
        
        # Create file object
        file_obj = io.BytesIO(document_content.encode('utf-8'))
        
        bucket_name = os.getenv("GCS_BUCKET_NAME")
        if not bucket_name:
            raise Exception("GCS_BUCKET_NAME environment variable is not set")
        
        # Upload to Google Cloud Storage
        public_url = upload_to_gcs(file_obj, file_name, "text/plain", bucket_name)
        return public_url
    except Exception as e:
        print(f"Error saving document to GCS: {e}")
        raise e

async def save_docx_to_gcs(document_content: str, document_type: str, generated_for: str, logo_url: str = None, content_type: str = "text"):
    """Save document to Google Cloud Storage as a .docx file with optional logo on each page"""
    try:
        # Create filename
        safe_name = "".join(c for c in generated_for if c.isalnum() or c in (' ', '-', '_')).strip()
        file_name = f"{document_type.replace(' ', '_')}_{safe_name.replace(' ', '_')[:30]}.docx"
        
        # Create Word document
        doc = Document()
        
        # Add logo to header if logo_url is provided
        if logo_url:
            await add_logo_to_header(doc, logo_url)
        
        # Process content based on type
        if content_type == "html":
            await process_html_content(doc, document_content)
        else:
            await process_text_content(doc, document_content)
        
        # Save document to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        bucket_name = os.getenv("GCS_BUCKET_NAME")
        if not bucket_name:
            raise Exception("GCS_BUCKET_NAME environment variable is not set")
        
        # Upload to Google Cloud Storage
        public_url = upload_to_gcs(doc_bytes, file_name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", bucket_name)
        return public_url
    except Exception as e:
        print(f"Error saving .docx document to GCS: {e}")
        raise e

async def process_html_content(doc, html_content):
    """Process HTML content and add to Word document with proper formatting"""
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Process the body content
        body = soup.find('body')
        if body:
            process_html_elements(doc, body)
        else:
            # If no body tag, process the whole content
            process_html_elements(doc, soup)
            
    except Exception as e:
        print(f"Error processing HTML content: {e}")
        # Fallback to text processing by extracting plain text from HTML
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            plain_text = soup.get_text()
            await process_text_content(doc, plain_text)
        except Exception as fallback_error:
            print(f"Error in fallback text processing: {fallback_error}")
            # Ultimate fallback - add raw content
            doc.add_paragraph("Error processing document content. Please contact support.")
            doc.add_paragraph(html_content[:500] + "..." if len(html_content) > 500 else html_content)

def process_html_elements(doc, element):
    """Recursively process HTML elements and add to Word document"""
    # Check if element has children attribute (i.e., it's a Tag, not NavigableString)
    if not hasattr(element, 'children'):
        return
    
    for child in element.children:
        # Skip empty text nodes
        if hasattr(child, 'string') and child.string and child.string.strip():
            # This is a text node with content - add as paragraph
            text = child.string.strip()
            if text and text not in ['\n', '\r', '\t', ' ']:
                paragraph = doc.add_paragraph(text)
                if paragraph.runs:
                    paragraph_run = paragraph.runs[0]
                    paragraph_run.font.size = Pt(11)
        elif hasattr(child, 'name') and child.name:
            # This is an HTML tag element
            if child.name == 'div' and 'title' in child.get('class', []):
                # Title
                text = child.get_text().strip()
                if text:
                    title_paragraph = doc.add_paragraph(text)
                    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if title_paragraph.runs:
                        title_run = title_paragraph.runs[0]
                        title_run.font.size = Pt(18)
                        title_run.font.bold = True
            elif child.name == 'div' and 'section-title' in child.get('class', []):
                # Section title
                text = child.get_text().strip()
                if text:
                    section_paragraph = doc.add_paragraph(text)
                    if section_paragraph.runs:
                        section_run = section_paragraph.runs[0]
                        section_run.font.size = Pt(14)
                        section_run.font.bold = True
            elif child.name == 'div' and 'parties' in child.get('class', []):
                # Parties section with background
                text = child.get_text().strip()
                if text:
                    parties_paragraph = doc.add_paragraph(text)
                    if parties_paragraph.runs:
                        parties_run = parties_paragraph.runs[0]
                        parties_run.font.size = Pt(11)
            elif child.name == 'div' and 'bullet-point' in child.get('class', []):
                # Bullet point
                text = child.get_text().strip()
                if text:
                    bullet_paragraph = doc.add_paragraph(text, style='List Bullet')
                    if bullet_paragraph.runs:
                        bullet_run = bullet_paragraph.runs[0]
                        bullet_run.font.size = Pt(11)
            elif child.name == 'div' and 'signature-section' in child.get('class', []):
                # Signature section
                doc.add_page_break()
                text = child.get_text().strip()
                if text:
                    sig_paragraph = doc.add_paragraph(text)
                    if sig_paragraph.runs:
                        sig_run = sig_paragraph.runs[0]
                        sig_run.font.size = Pt(11)
            elif child.name in ['div', 'p']:
                # Regular paragraph
                text = child.get_text().strip()
                if text:
                    paragraph = doc.add_paragraph(text)
                    if paragraph.runs:
                        paragraph_run = paragraph.runs[0]
                        paragraph_run.font.size = Pt(11)
            elif child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Headers
                text = child.get_text().strip()
                if text:
                    header_paragraph = doc.add_paragraph(text)
                    if header_paragraph.runs:
                        header_run = header_paragraph.runs[0]
                        header_run.font.size = Pt(16 - int(child.name[1]) * 2)
                        header_run.font.bold = True
            elif child.name in ['strong', 'b']:
                # Bold text
                text = child.get_text().strip()
                if text:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(text)
                    run.font.bold = True
                    run.font.size = Pt(11)
            elif child.name == 'br':
                # Line break
                doc.add_paragraph()
            else:
                # Process children recursively for other elements
                process_html_elements(doc, child)

async def process_text_content(doc, document_content):
    """Process plain text content and add to Word document"""
    # Split content into paragraphs and add to document
    paragraphs = document_content.split('\n\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            # Check if it's a title (all caps or contains specific keywords)
            if (paragraph.strip().isupper() and len(paragraph.strip()) < 100) or \
               any(keyword in paragraph.upper() for keyword in ['AGREEMENT', 'CONTRACT', 'PROPOSAL', 'POLICY', 'TERMS']):
                # Add as title
                title_paragraph = doc.add_paragraph(paragraph.strip())
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_paragraph.runs[0]
                title_run.font.size = Pt(16)
                title_run.font.bold = True
            elif paragraph.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                # Add as section header
                section_paragraph = doc.add_paragraph(paragraph.strip())
                section_run = section_paragraph.runs[0]
                section_run.font.size = Pt(12)
                section_run.font.bold = True
            else:
                # Add as regular paragraph
                doc.add_paragraph(paragraph.strip())

async def add_logo_to_header(doc, logo_url):
    """Add logo to document header"""
    try:
        # Get the header
        section = doc.sections[0]
        header = section.header
        
        # Create header paragraph
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Download logo image
        response = requests.get(logo_url)
        if response.status_code == 200:
            # Create temporary file for logo
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                tmp_file.write(response.content)
                tmp_file_path = tmp_file.name
            
            try:
                # Add logo to header
                run = header_paragraph.runs[0] if header_paragraph.runs else header_paragraph.add_run()
                run.add_picture(tmp_file_path, width=Inches(1.0))
                
                # Add company name next to logo
                run.add_text("  ScaleBuild AI")
                run.font.size = Pt(12)
                run.font.bold = True
                
            finally:
                # Clean up temp file
                os.unlink(tmp_file_path)
        else:
            # If logo can't be downloaded, just add text
            run = header_paragraph.runs[0] if header_paragraph.runs else header_paragraph.add_run()
            run.add_text("ScaleBuild AI")
            run.font.size = Pt(12)
            run.font.bold = True
            
    except Exception as e:
        print(f"Error adding logo to header: {e}")
        # If logo fails, just add text header
        header_paragraph = doc.sections[0].header.paragraphs[0]
        run = header_paragraph.runs[0] if header_paragraph.runs else header_paragraph.add_run()
        run.add_text("ScaleBuild AI")
        run.font.size = Pt(12)
        run.font.bold = True







