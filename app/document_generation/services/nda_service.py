from services.document_utils import save_docx_to_gcs
from datetime import datetime
from google.cloud import storage
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
import io
import uuid
import requests
from tempfile import NamedTemporaryFile

# --- NDA HTML Template ---
nda_html_template = """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Non-Disclosure Agreement</title>
    <style>
        body {{ font-family: 'Times New Roman', serif; font-size: 12pt; line-height: 1.5; }}
        .header {{ text-align: center; margin-bottom: 20px; }}
        .title {{ font-size: 14pt; font-weight: bold; text-transform: uppercase; }}
        .parties {{ margin: 20px 0; }}
        .party {{ font-weight: bold; margin: 5px 0; }}
        .section {{ margin: 15px 0; }}
        .section-title {{ font-weight: bold; }}
        .content {{ margin: 10px 0; }}
        .bullet-list {{ margin: 10px 0; padding-left: 20px; }}
        .signature {{ margin-top: 30px; }}
    </style>
</head>
<body>
    <div class="header">
        <div class="title">NON-DISCLOSURE AGREEMENT (NDA)</div>
        <p>This Agreement is made and entered into on this {effective_date}, by and between:</p>
    </div>
    
    <div class="parties">
        <div class="party">{disclosing_party}</div>
        <div>(hereinafter referred to as the "Disclosing Party" or "Company")</div>
        <br>
        <div style="text-align: center; font-weight: bold;">AND</div>
        <br>
        <div class="party">{receiving_party}</div>
        <div>(hereinafter referred to as the "Receiving Party" or "Employee")</div>
    </div>

    <div class="section">
        <div class="section-title">1. Purpose</div>
        <div class="content">
            This Agreement is entered into for the purpose of {purpose}. The Disclosing Party may disclose confidential and proprietary information to the Receiving Party which {receiving_party} may gain access in the course of her work via {disclosing_party}.
        </div>
    </div>

    <div class="section">
        <div class="section-title">2. Definition of Confidential Information</div>
        <div class="content">
            For purposes of this Agreement, "Confidential Information" shall include, but is not limited to:
            <ul class="bullet-list">
                <li>Project specifications, business strategies, technical documentation, source code, APIs, internal chatbot interactions, chat histories, and communications.</li>
                <li>Access credentials and usage of {disclosing_party}'s internal drives, servers, LLM APIs, tools, or environments.</li>
                <li>All material and non-public information shared by {disclosing_party} in the course of its work with {receiving_party}.</li>
                <li>Any data or insights obtained through observation or interaction with {disclosing_party}'s internal systems.</li>
            </ul>
            <p>Confidential Information may be in oral, written, digital, visual, or any other form, and shall remain protected regardless of the format or mode of delivery.</p>
        </div>
    </div>

    <div class="section">
        <div class="section-title">3. Obligations of the Receiving Party</div>
        <div class="content">
            The Receiving Party agrees that:
            <ul class="bullet-list">
                <li>She shall hold all Confidential Information in the strictest confidence and shall not disclose, share, or discuss it with anyone, including but not limited to employees, contractors, or management, without express prior written consent from the Disclosing Party.</li>
                <li>She shall use the Confidential Information solely for the purpose of performing assigned work in connection with the Disclosing Party's projects.</li>
                <li>She shall take all reasonable steps to protect the confidentiality of such information and prevent unauthorized use or disclosure.</li>
                <li>She shall immediately report any known or suspected breach of this Agreement to the Disclosing Party.</li>
            </ul>
        </div>
    </div>

    <div class="section">
        <div class="section-title">4. Exclusions</div>
        <div class="content">
            Confidential Information does not include information that:
            <ul class="bullet-list">
                <li>Is publicly known through no fault or breach of the Receiving Party;</li>
                <li>Is lawfully obtained by the Receiving Party from a third party without obligation of confidentiality;</li>
                <li>Is independently developed without use of or reference to the Disclosing Party's Confidential Information.</li>
            </ul>
        </div>
    </div>

    <div class="section">
        <div class="section-title">5. Term</div>
        <div class="content">
            This Agreement shall remain in effect for a period of {duration} and shall continue to bind the Receiving Party following the termination of that engagement.
        </div>
    </div>

    <div class="section">
        <div class="section-title">6. Return or Destruction of Information</div>
        <div class="content">
            Upon completion of work or upon request by the Disclosing Party, the Receiving Party agrees to promptly return or destroy all materials containing Confidential Information, including digital copies.
        </div>
    </div>

    <div class="section">
        <div class="section-title">7. Remedies</div>
        <div class="content">
            The Receiving Party acknowledges that any breach of this Agreement may cause irreparable harm to the Disclosing Party, for which monetary damages may be inadequate. Therefore, the Disclosing Party shall be entitled to seek equitable relief, including injunction and specific performance, in addition to all other remedies available at law or in equity.
        </div>
    </div>

    <div class="section">
        <div class="section-title">8. Governing Law</div>
        <div class="content">
            This Agreement shall be governed by and construed in accordance with the laws of {governing_law}, without regard to its conflict of laws principles.
        </div>
    </div>

    <div class="signature">
        <div class="section-title">9. Signatures</div>
        <div class="content">
            By signing below, both parties acknowledge that they have read, understood, and agree to be bound by the terms and conditions of this Agreement.
        </div>
        <br><br>
        <table width="100%">
            <tr>
                <td width="50%">
                    <strong>Disclosing Party:</strong> {disclosing_party}<br>
                    <br>_________________________<br>
                    Date: _______________
                </td>
                <td width="50%">
                    <strong>Receiving Party:</strong> {receiving_party}<br>
                    <br>_________________________<br>
                    Date: _______________
                </td>
            </tr>
        </table>
    </div>
</body>
</html>"""

# --- NDA Markdown Template ---
nda_markdown_template = """# NON-DISCLOSURE AGREEMENT (NDA)

**This Agreement is made and entered into on this {effective_date}, by and between:**

**{disclosing_party}**  
(hereinafter referred to as the "Disclosing Party" or "Company")

**AND**

**{receiving_party}**  
(hereinafter referred to as the "Receiving Party" or "Employee")

## 1. Purpose

This Agreement is entered into for the purpose of {purpose}. The Disclosing Party may disclose confidential and proprietary information to the Receiving Party which {receiving_party} may gain access in the course of her work via {disclosing_party}.

## 2. Definition of Confidential Information

For purposes of this Agreement, "Confidential Information" shall include, but is not limited to:

• Project specifications, business strategies, technical documentation, source code, APIs, internal chatbot interactions, chat histories, and communications.
• Access credentials and usage of {disclosing_party}'s internal drives, servers, LLM APIs, tools, or environments.
• All material and non-public information shared by {disclosing_party} in the course of its work with {receiving_party}.
• Any data or insights obtained through observation or interaction with {disclosing_party}'s internal systems.

Confidential Information may be in oral, written, digital, visual, or any other form, and shall remain protected regardless of the format or mode of delivery.

## 3. Obligations of the Receiving Party

The Receiving Party agrees that:

• She shall hold all Confidential Information in the strictest confidence and shall not disclose, share, or discuss it with anyone, including but not limited to employees, contractors, or management, without express prior written consent from the Disclosing Party.
• She shall use the Confidential Information solely for the purpose of performing assigned work in connection with the Disclosing Party's projects.
• She shall take all reasonable steps to protect the confidentiality of such information and prevent unauthorized use or disclosure.
• She shall immediately report any known or suspected breach of this Agreement to the Disclosing Party.

## 4. Exclusions

Confidential Information does not include information that:

• Is publicly known through no fault or breach of the Receiving Party;
• Is lawfully obtained by the Receiving Party from a third party without obligation of confidentiality;
• Is independently developed without use of or reference to the Disclosing Party's Confidential Information.

## 5. Term

This Agreement shall remain in effect for a period of {duration} and shall continue to bind the Receiving Party following the termination of that engagement.

## 6. Return or Destruction of Information

Upon completion of work or upon request by the Disclosing Party, the Receiving Party agrees to promptly return or destroy all materials containing Confidential Information, including digital copies.

## 7. Remedies

The Receiving Party acknowledges that any breach of this Agreement may cause irreparable harm to the Disclosing Party, for which monetary damages may be inadequate. Therefore, the Disclosing Party shall be entitled to seek equitable relief, including injunction and specific performance, in addition to all other remedies available at law or in equity.

## 8. Governing Law

This Agreement shall be governed by and construed in accordance with the laws of {governing_law}, without regard to its conflict of laws principles.

## 9. Signatures

By signing below, both parties acknowledge that they have read, understood, and agree to be bound by the terms and conditions of this Agreement.

**Disclosing Party:** {disclosing_party}

_________________________  
Date: _______________

**Receiving Party:** {receiving_party}

_________________________  
Date: _______________

---

*This document was generated on {effective_date}*
*Non-Disclosure Agreement - Confidential*
"""

async def upload_to_gcs(content: str, filename: str, content_type: str = "text/plain"):
    """Upload content to Google Cloud Storage and return the public URL"""
    try:
        # Initialize GCS client
        client = storage.Client()
        bucket_name = "deck123"  # Your bucket name
        bucket = client.bucket(bucket_name)
        
        # Create blob
        blob = bucket.blob(filename)
        
        # Upload content
        blob.upload_from_string(content, content_type=content_type)
        
        # Make blob public
        blob.make_public()
        
        return blob.public_url
    except Exception as e:
        print(f"Error uploading to GCS: {e}")
        return None

def add_logo_to_header(doc, logo_url):
    """Add logo image to the header of every page"""
    if not logo_url:
        print("No logo URL provided")
        return
        
    try:
        print(f"Attempting to download logo from: {logo_url}")
        
        # Download the logo from URL with proper headers
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(logo_url, headers=headers, timeout=15, stream=True)
        response.raise_for_status()
        
        print(f"Logo downloaded successfully. Status: {response.status_code}")
        print(f"Content-Type: {response.headers.get('Content-Type', 'Unknown')}")
        print(f"Logo size: {len(response.content)} bytes")
        
        # Validate content size
        if len(response.content) == 0:
            raise ValueError("Downloaded logo file is empty")
            
        # Determine file extension from URL or content type
        if logo_url.lower().endswith('.png'):
            suffix = '.png'
        elif logo_url.lower().endswith(('.jpg', '.jpeg')):
            suffix = '.jpg'
        elif logo_url.lower().endswith('.gif'):
            suffix = '.gif'
        elif logo_url.lower().endswith('.webp'):
            suffix = '.webp'
        else:
            content_type = response.headers.get('Content-Type', '').lower()
            if 'png' in content_type:
                suffix = '.png'
            elif 'jpeg' in content_type or 'jpg' in content_type:
                suffix = '.jpg'
            elif 'gif' in content_type:
                suffix = '.gif'
            elif 'webp' in content_type:
                suffix = '.webp'
            else:
                suffix = '.png'  # Default to PNG
        
        # Create a temporary file to store the logo
        with NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(response.content)
            tmp_file_path = tmp_file.name
        
        print(f"Logo saved to temporary file: {tmp_file_path}")
        
        # Add logo to header of each section
        for section in doc.sections:
            header = section.header
            
            # Clear existing header paragraphs and create new one
            for paragraph in header.paragraphs:
                paragraph.clear()
            
            # If no paragraphs exist, add one
            if not header.paragraphs:
                header_p = header.add_paragraph()
            else:
                header_p = header.paragraphs[0]
            
            header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add logo image
            run = header_p.add_run()
            try:
                run.add_picture(tmp_file_path, width=Inches(1.0), height=Inches(0.6))
                print(f"Logo successfully added to header for section")
            except Exception as pic_error:
                print(f"Error adding picture to run: {pic_error}")
                raise pic_error
        
        # Clean up temporary file
        os.unlink(tmp_file_path)
        print("Logo added to header successfully and temp file cleaned up")
        
    except requests.exceptions.RequestException as e:
        print(f"Error downloading logo: {e}")
        # Don't add fallback - let the document be generated without logo
        pass
    except Exception as e:
        print(f"Error adding logo to header: {e}")
        import traceback
        traceback.print_exc()
        # Don't add fallback - let the document be generated without logo
        pass

def create_professional_docx(data: dict):
    """Create a professional DOCX document using python-docx"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_heading("NON-DISCLOSURE AGREEMENT (NDA)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(14)
    title_run.font.name = "Times New Roman"
    title_run.bold = True

    # Agreement intro
    intro = doc.add_paragraph(f"This Agreement is made and entered into on this {data.get('effective_date', '')}, by and between:")
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Disclosing Party
    dp = doc.add_paragraph()
    dp_run = dp.add_run(data.get('disclosing_party', ''))
    dp_run.bold = True
    dp.add_run('\n(hereinafter referred to as the "Disclosing Party" or "Company")')

    # AND separator
    and_p = doc.add_paragraph("AND")
    and_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    and_p.runs[0].bold = True

    # Receiving Party
    rp = doc.add_paragraph()
    rp_run = rp.add_run(data.get('receiving_party', ''))
    rp_run.bold = True
    rp.add_run('\n(hereinafter referred to as the "Receiving Party" or "Employee")')

    # Sections
    sections_data = [
        ("1. Purpose", f"This Agreement is entered into for the purpose of {data.get('purpose', '')}. The Disclosing Party may disclose confidential and proprietary information to the Receiving Party which {data.get('receiving_party', '')} may gain access in the course of her work via {data.get('disclosing_party', '')}."),
        ("2. Definition of Confidential Information", "For purposes of this Agreement, \"Confidential Information\" shall include, but is not limited to:"),
        ("3. Obligations of the Receiving Party", "The Receiving Party agrees that:"),
        ("4. Exclusions", "Confidential Information does not include information that:"),
        ("5. Term", f"This Agreement shall remain in effect for a period of {data.get('duration', '')} and shall continue to bind the Receiving Party following the termination of that engagement."),
        ("6. Return or Destruction of Information", "Upon completion of work or upon request by the Disclosing Party, the Receiving Party agrees to promptly return or destroy all materials containing Confidential Information, including digital copies."),
        ("7. Remedies", "The Receiving Party acknowledges that any breach of this Agreement may cause irreparable harm to the Disclosing Party, for which monetary damages may be inadequate. Therefore, the Disclosing Party shall be entitled to seek equitable relief, including injunction and specific performance, in addition to all other remedies available at law or in equity."),
        ("8. Governing Law", f"This Agreement shall be governed by and construed in accordance with the laws of {data.get('governing_law', '')}, without regard to its conflict of laws principles."),
        ("9. Signatures", "By signing below, both parties acknowledge that they have read, understood, and agree to be bound by the terms and conditions of this Agreement.")
    ]

    for section_title, section_content in sections_data:
        # Section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section_title)
        heading_run.bold = True
        heading_run.font.name = "Times New Roman"
        heading_run.font.size = Pt(12)

        # Section content
        content = doc.add_paragraph(section_content)
        content_run = content.runs[0]
        content_run.font.name = "Times New Roman"
        content_run.font.size = Pt(12)

        # Add bullet points for specific sections
        if section_title == "2. Definition of Confidential Information":
            bullets = [
                f"Project specifications, business strategies, technical documentation, source code, APIs, internal chatbot interactions, chat histories, and communications.",
                f"Access credentials and usage of {data.get('disclosing_party', '')}'s internal drives, servers, LLM APIs, tools, or environments.",
                f"All material and non-public information shared by {data.get('disclosing_party', '')} in the course of its work with {data.get('receiving_party', '')}.",
                f"Any data or insights obtained through observation or interaction with {data.get('disclosing_party', '')}'s internal systems."
            ]
            for bullet in bullets:
                bullet_p = doc.add_paragraph(bullet, style='List Bullet')
                bullet_p.runs[0].font.name = "Times New Roman"
                bullet_p.runs[0].font.size = Pt(12)

            doc.add_paragraph("Confidential Information may be in oral, written, digital, visual, or any other form, and shall remain protected regardless of the format or mode of delivery.")

        elif section_title == "3. Obligations of the Receiving Party":
            bullets = [
                "She shall hold all Confidential Information in the strictest confidence and shall not disclose, share, or discuss it with anyone, including but not limited to employees, contractors, or management, without express prior written consent from the Disclosing Party.",
                "She shall use the Confidential Information solely for the purpose of performing assigned work in connection with the Disclosing Party's projects.",
                "She shall take all reasonable steps to protect the confidentiality of such information and prevent unauthorized use or disclosure.",
                "She shall immediately report any known or suspected breach of this Agreement to the Disclosing Party."
            ]
            for bullet in bullets:
                bullet_p = doc.add_paragraph(bullet, style='List Bullet')
                bullet_p.runs[0].font.name = "Times New Roman"
                bullet_p.runs[0].font.size = Pt(12)

        elif section_title == "4. Exclusions":
            bullets = [
                "Is publicly known through no fault or breach of the Receiving Party;",
                "Is lawfully obtained by the Receiving Party from a third party without obligation of confidentiality;",
                "Is independently developed without use of or reference to the Disclosing Party's Confidential Information."
            ]
            for bullet in bullets:
                bullet_p = doc.add_paragraph(bullet, style='List Bullet')
                bullet_p.runs[0].font.name = "Times New Roman"
                bullet_p.runs[0].font.size = Pt(12)

    # Signature section
    doc.add_paragraph("\n")
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = 'Table Grid'

    # Headers
    sig_table.cell(0, 0).text = f"Disclosing Party: {data.get('disclosing_party', '')}"
    sig_table.cell(0, 1).text = f"Receiving Party: {data.get('receiving_party', '')}"

    # Signature lines
    sig_table.cell(1, 0).text = "\n_________________________"
    sig_table.cell(1, 1).text = "\n_________________________"

    # Date lines
    sig_table.cell(2, 0).text = "Date: _______________"
    sig_table.cell(2, 1).text = "Date: _______________"

    return doc

def add_page_number(doc):
    """Add page number to the document"""
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.text = "1"

async def generate_nda(data: dict):
    """Generate an NDA document with HTML, Markdown, and DOCX formats"""
    try:
        # Generate HTML content
        html_content = nda_html_template.format(
            disclosing_party=data.get("disclosing_party", ""),
            receiving_party=data.get("receiving_party", ""),
            purpose=data.get("purpose", ""),
            duration=data.get("duration", ""),
            governing_law=data.get("governing_law", ""),
            effective_date=data.get("effective_date", "")
        )

        # Generate Markdown content
        markdown_content = nda_markdown_template.format(
            disclosing_party=data.get("disclosing_party", ""),
            receiving_party=data.get("receiving_party", ""),
            purpose=data.get("purpose", ""),
            duration=data.get("duration", ""),
            governing_law=data.get("governing_law", ""),
            effective_date=data.get("effective_date", "")
        )

        # Create professional DOCX
        doc = create_professional_docx(data)

        # Add logo to header using logo URL from request body
        logo_url = data.get("logo_url")
        if logo_url:
            print(f"Logo URL found in request: {logo_url}")
            add_logo_to_header(doc, logo_url)
        else:
            print("No logo URL provided in request data")

        # Save DOCX to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        # Generate unique filenames
        unique_id = str(uuid.uuid4())[:8]
        base_filename = f"NDA_{data.get('disclosing_party', '').replace(' ', '_')}_{data.get('receiving_party', '').replace(' ', '_')}_{unique_id}"

        # Upload files to GCS
        html_url = await upload_to_gcs(html_content, f"{base_filename}.html", "text/html")
        markdown_url = await upload_to_gcs(markdown_content, f"{base_filename}.md", "text/markdown")

        # Upload DOCX to GCS
        client = storage.Client()
        bucket = client.bucket("deck123")
        docx_blob = bucket.blob(f"{base_filename}.docx")
        docx_blob.upload_from_string(docx_bytes.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        docx_blob.make_public()
        docx_url = docx_blob.public_url

        return {
            "document_content": html_content,
            "markdown_content": markdown_content,
            "document_type": "Non-Disclosure Agreement",
            "generated_for": f"{data.get('disclosing_party')} & {data.get('receiving_party')}",
            "creation_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "word_count": len(html_content.split()),
            "format": "Multiple formats (HTML, Markdown, DOCX)",
            "urls": {
                "html": html_url,
                "markdown": markdown_url,
                "docx": docx_url
            },
            "document_url": docx_url  # Main document URL for compatibility
        }
    except Exception as e:
        print(f"Error in NDA generation: {e}")
        raise e
