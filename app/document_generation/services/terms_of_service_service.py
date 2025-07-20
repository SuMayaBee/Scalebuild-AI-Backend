from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.schema.output_parser import StrOutputParser
from services.document_utils import save_docx_to_gcs, upload_to_gcs
from datetime import datetime
import os
import tempfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.enum.section import WD_SECTION
import requests
import re
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- OpenAI Model ---
model = ChatOpenAI(model_name="gpt-4o", temperature=0.3)

# --- Terms of Service Template ---
terms_of_service_template = """You are a legal document specialist. Create comprehensive Terms of Service based on the following information:

Company Name: {company_name}
Website URL: {website_url}
Company Address: {company_address}
Service Description: {service_description}
User Responsibilities: {user_responsibilities}
Prohibited Activities: {prohibited_activities}
Payment Terms: {payment_terms}
Cancellation Policy: {cancellation_policy}
Limitation of Liability: {limitation_of_liability}
Governing Law: {governing_law}
Contact Email: {contact_email}

Generate comprehensive Terms of Service that include:

1. Acceptance of Terms
2. Description of Service
3. User Accounts and Registration
4. User Responsibilities and Conduct
5. Prohibited Uses
6. Payment Terms and Billing
7. Cancellation and Refunds
8. Intellectual Property Rights
9. Privacy and Data Protection
10. Limitation of Liability
11. Indemnification
12. Termination
13. Governing Law and Disputes
14. Changes to Terms
15. Contact Information

Use clear, legally sound language appropriate for online services."""

terms_of_service_prompt = PromptTemplate.from_template(terms_of_service_template)
terms_of_service_chain = terms_of_service_prompt | model | StrOutputParser()

def add_logo_and_page_number_to_header(doc, logo_url):
    """Add logo on left and page number on right in header"""
    if not logo_url:
        logger.warning("No logo URL provided")
        return
        
    try:
        logger.info(f"Attempting to download logo from: {logo_url}")
        
        # Download the logo from URL with proper headers
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(logo_url, headers=headers, timeout=15, stream=True)
        response.raise_for_status()
        
        logger.info(f"Logo downloaded successfully. Status: {response.status_code}")
        logger.info(f"Content-Type: {response.headers.get('Content-Type', 'Unknown')}")
        logger.info(f"Logo size: {len(response.content)} bytes")
        
        # Validate content size
        if len(response.content) == 0:
            raise ValueError("Downloaded logo file is empty")
            
        # Determine file extension from URL or content type
        if logo_url.lower().endswith('.png'):
            suffix = '.png'
        elif logo_url.lower().endswith(('.jpg', '.jpeg')):
            suffix = '.jpg'
        elif logo_url.lower().endswith('.gif'):
            suffix = '.gif'
        elif logo_url.lower().endswith('.webp'):
            suffix = '.webp'
        else:
            content_type = response.headers.get('Content-Type', '').lower()
            if 'png' in content_type:
                suffix = '.png'
            elif 'jpeg' in content_type or 'jpg' in content_type:
                suffix = '.jpg'
            elif 'gif' in content_type:
                suffix = '.gif'
            elif 'webp' in content_type:
                suffix = '.webp'
            else:
                suffix = '.png'  # Default to PNG
        
        # Create a temporary file to store the logo
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(response.content)
            tmp_file_path = tmp_file.name
        
        logger.info(f"Logo saved to temporary file: {tmp_file_path}")
        
        # Add logo and page number to header of each section
        for section in doc.sections:
            header = section.header
            
            # Clear existing header paragraphs and create new one
            for paragraph in header.paragraphs:
                paragraph.clear()
            
            # If no paragraphs exist, add one
            if not header.paragraphs:
                header_p = header.add_paragraph()
            else:
                header_p = header.paragraphs[0]
            
            # Create a table for layout (logo left, page number right)
            from docx import Document
            header_table = header_p._element.getparent().add_table(rows=1, cols=2)
            header_table.width = Inches(6.5)
            
            # Left cell for logo
            left_cell = header_table.cell(0, 0)
            left_cell.width = Inches(2)
            left_para = left_cell.paragraphs[0]
            left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add logo image
            try:
                left_run = left_para.add_run()
                left_run.add_picture(tmp_file_path, width=Inches(1.0), height=Inches(0.6))
                logger.info("Logo successfully added to header")
            except Exception as pic_error:
                logger.error(f"Error adding picture to run: {pic_error}")
                raise pic_error
            
            # Right cell for page number
            right_cell = header_table.cell(0, 1)
            right_cell.width = Inches(4.5)
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add page number field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            page_run = right_para.add_run()
            page_run._r.append(fldChar1)
            page_run._r.append(instrText)
            page_run._r.append(fldChar2)
            page_run.font.name = "Times New Roman"
            page_run.font.size = Pt(10)
            
            # Remove table borders
            for row in header_table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = tcPr.first_child_found_in("w:tcBorders")
                    if tcBorders is None:
                        tcBorders = OxmlElement('w:tcBorders')
                        tcPr.append(tcBorders)
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = tcBorders.first_child_found_in(f"w:{border_name}")
                        if border is None:
                            border = OxmlElement(f'w:{border_name}')
                            tcBorders.append(border)
                        border.set(qn('w:val'), 'nil')
        
        # Clean up temporary file
        os.unlink(tmp_file_path)
        logger.info("Logo added to header successfully and temp file cleaned up")
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Error downloading logo: {e}")
        # Don't add fallback - let the document be generated without logo
        pass
    except Exception as e:
        logger.error(f"Error adding logo to header: {e}")
        import traceback
        traceback.print_exc()
        # Don't add fallback - let the document be generated without logo
        pass

def create_professional_docx(content: str, title: str):
    """Create a professional DOCX document using python-docx"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.5)  # Extra space for header
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title_paragraph = doc.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_run.bold = True

    # Process content line by line
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if line is a section header (starts with number or ##)
        if (line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.', '11.', '12.', '13.', '14.', '15.')) or 
            line.startswith('##')):
            # Section heading
            heading = doc.add_paragraph()
            heading_text = line.replace('##', '').strip()
            heading_run = heading.add_run(heading_text)
            heading_run.bold = True
            heading_run.font.name = "Times New Roman"
            heading_run.font.size = Pt(14)
        else:
            # Regular paragraph with markdown formatting
            paragraph = doc.add_paragraph()
            parse_markdown_to_runs(paragraph, line)

    return doc

def parse_markdown_to_runs(paragraph, text):
    """Parse markdown formatting in text and add formatted runs to paragraph"""
    
    # Split text by ** markers to identify bold sections
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text - remove ** markers
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        else:
            # Regular text
            if part:  # Only add non-empty parts
                run = paragraph.add_run(part)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

# --- Terms of Service Chain ---
async def generate_terms_of_service(data: dict):
    """Generate Terms of Service document using GPT-4o and save as .docx with logo and page numbers"""
    try:
        from google.cloud import storage
        import uuid
        import io
        
        user_responsibilities_list = ", ".join(data.get("user_responsibilities", []))
        prohibited_activities_list = ", ".join(data.get("prohibited_activities", []))
        
        document_content = ""
        async for chunk in terms_of_service_chain.astream({
            "company_name": data.get("company_name"),
            "website_url": data.get("website_url"),
            "company_address": data.get("company_address"),
            "service_description": data.get("service_description"),
            "user_responsibilities": user_responsibilities_list,
            "prohibited_activities": prohibited_activities_list,
            "payment_terms": data.get("payment_terms"),
            "cancellation_policy": data.get("cancellation_policy"),
            "limitation_of_liability": data.get("limitation_of_liability"),
            "governing_law": data.get("governing_law"),
            "contact_email": data.get("contact_email")
        }):
            document_content += chunk
        
        # Create professional DOCX
        doc = create_professional_docx(document_content, "Terms of Service")

        # Add logo and page numbers to header using logo URL from request body
        logo_url = data.get("logo_url")
        if logo_url:
            logger.info(f"Logo URL found in request: {logo_url}")
            add_logo_and_page_number_to_header(doc, logo_url)
        else:
            logger.warning("No logo URL provided in request data")

        # Save DOCX to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        # Generate unique filename
        unique_id = str(uuid.uuid4())[:8]
        filename = f"Terms_of_Service_{data.get('company_name', '').replace(' ', '_')}_{unique_id}.docx"

        # Upload DOCX to GCS
        client = storage.Client()
        bucket = client.bucket("deck123")
        docx_blob = bucket.blob(filename)
        docx_blob.upload_from_string(docx_bytes.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        docx_blob.make_public()
        document_url = docx_blob.public_url
        
        return {
            "document_content": document_content.strip(),
            "document_url": document_url,
            "document_type": "Terms of Service",
            "generated_for": data.get("company_name"),
            "creation_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "word_count": len(document_content.split()),
            "format": "DOCX with logo and page numbers"
        }
    except Exception as e:
        logger.error(f"Error in Terms of Service generation: {e}")
        raise e
