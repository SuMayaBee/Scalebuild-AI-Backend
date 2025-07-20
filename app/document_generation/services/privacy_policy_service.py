from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.schema.output_parser import StrOutputParser
from services.document_utils import save_docx_to_gcs
from datetime import datetime
from google.cloud import storage
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
import io
import uuid
import requests
from tempfile import NamedTemporaryFile

# --- OpenAI Model ---
model = ChatOpenAI(model_name="gpt-4o", temperature=0.3)

# --- Privacy Policy Template ---
privacy_policy_template = """You are a legal document specialist. Create a comprehensive Privacy Policy based on the following information:

Company Name: {company_name}
Website URL: {website_url}
Company Address: {company_address}
Data Collected: {data_collected}
Data Usage Purpose: {data_usage_purpose}
Third Party Sharing: {third_party_sharing}
Data Retention Period: {data_retention_period}
User Rights: {user_rights}
Cookies Usage: {cookies_usage}
Contact Email: {contact_email}
Governing Law: {governing_law}
Effective Date: {effective_date}

Generate a comprehensive Privacy Policy that includes:

1. Information We Collect
2. How We Use Your Information
3. Information Sharing and Disclosure
4. Data Security
5. Data Retention
6. Your Rights and Choices
7. Cookies and Tracking Technologies
8. Children's Privacy
9. International Data Transfers
10. Changes to Privacy Policy
11. Contact Information
12. Compliance with Laws (GDPR, CCPA, etc.)

Ensure compliance with major privacy regulations and use clear, accessible language."""

privacy_policy_prompt = PromptTemplate.from_template(privacy_policy_template)
privacy_policy_chain = privacy_policy_prompt | model | StrOutputParser()

async def generate_privacy_policy(data: dict):
    """Generate Privacy Policy document using GPT-4o and save as .docx with logo"""
    try:
        data_collected_list = ", ".join(data.get("data_collected", []))
        data_usage_list = ", ".join(data.get("data_usage_purpose", []))
        user_rights_list = ", ".join(data.get("user_rights", []))
        
        document_content = ""
        async for chunk in privacy_policy_chain.astream({
            "company_name": data.get("company_name"),
            "website_url": data.get("website_url"),
            "company_address": data.get("company_address"),
            "data_collected": data_collected_list,
            "data_usage_purpose": data_usage_list,
            "third_party_sharing": data.get("third_party_sharing"),
            "data_retention_period": data.get("data_retention_period"),
            "user_rights": user_rights_list,
            "cookies_usage": data.get("cookies_usage"),
            "contact_email": data.get("contact_email"),
            "governing_law": data.get("governing_law"),
            "effective_date": data.get("effective_date")
        }):
            document_content += chunk
        
        # Create professional DOCX
        doc = create_professional_docx(document_content, "Privacy Policy")

        # Add logo to header using logo URL from request body
        logo_url = data.get("logo_url")
        if logo_url:
            print(f"Logo URL found in request: {logo_url}")
            add_logo_to_header(doc, logo_url)
        else:
            print("No logo URL provided in request data")

        # Add page numbers to footer
        add_page_numbers(doc)

        # Save DOCX to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        # Generate unique filename
        unique_id = str(uuid.uuid4())[:8]
        filename = f"Privacy_Policy_{data.get('company_name', '').replace(' ', '_')}_{unique_id}.docx"

        # Upload DOCX to GCS
        client = storage.Client()
        bucket = client.bucket("deck123")
        docx_blob = bucket.blob(filename)
        docx_blob.upload_from_string(docx_bytes.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        docx_blob.make_public()
        document_url = docx_blob.public_url
        
        return {
            "document_content": document_content.strip(),
            "document_url": document_url,
            "document_type": "Privacy Policy",
            "generated_for": data.get("company_name"),
            "creation_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "word_count": len(document_content.split()),
            "format": "DOCX with logo and page numbers"
        }
    except Exception as e:
        print(f"Error in Privacy Policy generation: {e}")
        raise e

async def upload_to_gcs(content: str, filename: str, content_type: str = "text/plain"):
    """Upload content to Google Cloud Storage and return the public URL"""
    try:
        # Initialize GCS client
        client = storage.Client()
        bucket_name = "deck123"  # Your bucket name
        bucket = client.bucket(bucket_name)
        
        # Create blob
        blob = bucket.blob(filename)
        
        # Upload content
        blob.upload_from_string(content, content_type=content_type)
        
        # Make blob public
        blob.make_public()
        
        return blob.public_url
    except Exception as e:
        print(f"Error uploading to GCS: {e}")
        return None

def add_logo_to_header(doc, logo_url):
    """Add logo image to the header of every page"""
    if not logo_url:
        print("No logo URL provided")
        return
        
    try:
        print(f"Attempting to download logo from: {logo_url}")
        
        # Download the logo from URL with proper headers
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(logo_url, headers=headers, timeout=15, stream=True)
        response.raise_for_status()
        
        print(f"Logo downloaded successfully. Status: {response.status_code}")
        print(f"Content-Type: {response.headers.get('Content-Type', 'Unknown')}")
        print(f"Logo size: {len(response.content)} bytes")
        
        # Validate content size
        if len(response.content) == 0:
            raise ValueError("Downloaded logo file is empty")
            
        # Determine file extension from URL or content type
        if logo_url.lower().endswith('.png'):
            suffix = '.png'
        elif logo_url.lower().endswith(('.jpg', '.jpeg')):
            suffix = '.jpg'
        elif logo_url.lower().endswith('.gif'):
            suffix = '.gif'
        elif logo_url.lower().endswith('.webp'):
            suffix = '.webp'
        else:
            content_type = response.headers.get('Content-Type', '').lower()
            if 'png' in content_type:
                suffix = '.png'
            elif 'jpeg' in content_type or 'jpg' in content_type:
                suffix = '.jpg'
            elif 'gif' in content_type:
                suffix = '.gif'
            elif 'webp' in content_type:
                suffix = '.webp'
            else:
                suffix = '.png'  # Default to PNG
        
        # Create a temporary file to store the logo
        with NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(response.content)
            tmp_file_path = tmp_file.name
        
        print(f"Logo saved to temporary file: {tmp_file_path}")
        
        # Add logo to header of each section
        for section in doc.sections:
            header = section.header
            
            # Clear existing header paragraphs and create new one
            for paragraph in header.paragraphs:
                paragraph.clear()
            
            # If no paragraphs exist, add one
            if not header.paragraphs:
                header_p = header.add_paragraph()
            else:
                header_p = header.paragraphs[0]
            
            header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add logo image
            run = header_p.add_run()
            try:
                run.add_picture(tmp_file_path, width=Inches(1.0), height=Inches(0.6))
                print(f"Logo successfully added to header for section")
            except Exception as pic_error:
                print(f"Error adding picture to run: {pic_error}")
                raise pic_error
        
        # Clean up temporary file
        os.unlink(tmp_file_path)
        print("Logo added to header successfully and temp file cleaned up")
        
    except requests.exceptions.RequestException as e:
        print(f"Error downloading logo: {e}")
        # Don't add fallback - let the document be generated without logo
        pass
    except Exception as e:
        print(f"Error adding logo to header: {e}")
        import traceback
        traceback.print_exc()
        # Don't add fallback - let the document be generated without logo
        pass

def add_page_numbers(doc):
    """Add page numbers to the document footer"""
    try:
        for section in doc.sections:
            footer = section.footer
            footer_p = footer.paragraphs[0]
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add page number field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run = footer_p.add_run()
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            
    except Exception as e:
        print(f"Error adding page numbers: {e}")

def create_professional_docx(content: str, title: str):
    """Create a professional DOCX document using python-docx"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title_paragraph = doc.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_run.bold = True

    # Process content line by line
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if line is a section header (starts with number or ##)
        if (line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.', '11.', '12.')) or 
            line.startswith('##')):
            # Section heading
            heading = doc.add_paragraph()
            heading_text = line.replace('##', '').strip()
            heading_run = heading.add_run(heading_text)
            heading_run.bold = True
            heading_run.font.name = "Times New Roman"
            heading_run.font.size = Pt(14)
        else:
            # Regular paragraph with markdown formatting
            paragraph = doc.add_paragraph()
            parse_markdown_to_runs(paragraph, line)

    return doc

def parse_markdown_to_runs(paragraph, text):
    """Parse markdown formatting in text and add formatted runs to paragraph"""
    import re
    
    # Split text by ** markers to identify bold sections
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text - remove ** markers
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        else:
            # Regular text
            if part:  # Only add non-empty parts
                run = paragraph.add_run(part)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
