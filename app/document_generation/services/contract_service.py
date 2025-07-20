from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.schema.output_parser import StrOutputParser
from services.document_utils import save_docx_to_gcs
from datetime import datetime
from google.cloud import storage
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
import io
import uuid
import requests
from tempfile import NamedTemporaryFile

# --- OpenAI Model ---
model = ChatOpenAI(model_name="gpt-4o", temperature=0.3)

# --- Contract Template ---
contract_template = """You are a legal document specialist. Create a comprehensive Contract based on the following information:

Contract Type: {contract_type}
Party 1 Name: {party1_name}
Party 1 Address: {party1_address}
Party 2 Name: {party2_name}
Party 2 Address: {party2_address}
Service Description: {service_description}
Contract Value: {contract_value}
Payment Terms: {payment_terms}
Duration: {duration}
Deliverables: {deliverables}
Terms & Conditions: {terms_conditions}
Effective Date: {effective_date}

Generate a formal Contract that includes:

1. Contract Overview and Scope
2. Parties Involved
3. Services/Deliverables
4. Payment Terms and Schedule
5. Timeline and Milestones
6. Terms and Conditions
7. Termination Clauses
8. Intellectual Property Rights
9. Liability and Indemnification
10. Dispute Resolution
11. Governing Law
12. Signatures

Use proper legal language and ensure the contract is comprehensive and enforceable."""

contract_prompt = PromptTemplate.from_template(contract_template)
contract_chain = contract_prompt | model | StrOutputParser()

async def upload_to_gcs(content: str, filename: str, content_type: str = "text/plain"):
    """Upload content to Google Cloud Storage and return the public URL"""
    try:
        # Initialize GCS client
        client = storage.Client()
        bucket_name = "deck123"  # Your bucket name
        bucket = client.bucket(bucket_name)
        
        # Create blob
        blob = bucket.blob(filename)
        
        # Upload content
        blob.upload_from_string(content, content_type=content_type)
        
        # Make blob public
        blob.make_public()
        
        return blob.public_url
    except Exception as e:
        print(f"Error uploading to GCS: {e}")
        return None

def add_logo_to_header(doc, logo_url):
    """Add logo image to the header of every page"""
    if not logo_url:
        print("No logo URL provided")
        return
        
    try:
        print(f"Attempting to download logo from: {logo_url}")
        
        # Download the logo from URL with proper headers
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(logo_url, headers=headers, timeout=15, stream=True)
        response.raise_for_status()
        
        print(f"Logo downloaded successfully. Status: {response.status_code}")
        print(f"Content-Type: {response.headers.get('Content-Type', 'Unknown')}")
        print(f"Logo size: {len(response.content)} bytes")
        
        # Validate content size
        if len(response.content) == 0:
            raise ValueError("Downloaded logo file is empty")
            
        # Determine file extension from URL or content type
        if logo_url.lower().endswith('.png'):
            suffix = '.png'
        elif logo_url.lower().endswith(('.jpg', '.jpeg')):
            suffix = '.jpg'
        elif logo_url.lower().endswith('.gif'):
            suffix = '.gif'
        elif logo_url.lower().endswith('.webp'):
            suffix = '.webp'
        else:
            content_type = response.headers.get('Content-Type', '').lower()
            if 'png' in content_type:
                suffix = '.png'
            elif 'jpeg' in content_type or 'jpg' in content_type:
                suffix = '.jpg'
            elif 'gif' in content_type:
                suffix = '.gif'
            elif 'webp' in content_type:
                suffix = '.webp'
            else:
                suffix = '.png'  # Default to PNG
        
        # Create a temporary file to store the logo
        with NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(response.content)
            tmp_file_path = tmp_file.name
        
        print(f"Logo saved to temporary file: {tmp_file_path}")
        
        # Add logo to header of each section
        for section in doc.sections:
            header = section.header
            
            # Clear existing header paragraphs and create new one
            for paragraph in header.paragraphs:
                paragraph.clear()
            
            # If no paragraphs exist, add one
            if not header.paragraphs:
                header_p = header.add_paragraph()
            else:
                header_p = header.paragraphs[0]
            
            header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add logo image
            run = header_p.add_run()
            try:
                run.add_picture(tmp_file_path, width=Inches(1.0), height=Inches(0.6))
                print(f"Logo successfully added to header for section")
            except Exception as pic_error:
                print(f"Error adding picture to run: {pic_error}")
                raise pic_error
        
        # Clean up temporary file
        os.unlink(tmp_file_path)
        print("Logo added to header successfully and temp file cleaned up")
        
    except requests.exceptions.RequestException as e:
        print(f"Error downloading logo: {e}")
        # Don't add fallback - let the document be generated without logo
        pass
    except Exception as e:
        print(f"Error adding logo to header: {e}")
        import traceback
        traceback.print_exc()
        # Don't add fallback - let the document be generated without logo
        pass

def create_professional_docx(content: str, title: str):
    """Create a professional DOCX document using python-docx"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title_paragraph = doc.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_run.bold = True

    # Process content line by line
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if line is a section header (starts with number or ##)
        if (line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.', '11.', '12.')) or 
            line.startswith('##')):
            # Section heading
            heading = doc.add_paragraph()
            heading_text = line.replace('##', '').strip()
            heading_run = heading.add_run(heading_text)
            heading_run.bold = True
            heading_run.font.name = "Times New Roman"
            heading_run.font.size = Pt(14)
        else:
            # Regular paragraph with markdown formatting
            paragraph = doc.add_paragraph()
            parse_markdown_to_runs(paragraph, line)

    return doc

def parse_markdown_to_runs(paragraph, text):
    """Parse markdown formatting in text and add formatted runs to paragraph"""
    import re
    
    # Split text by ** markers to identify bold sections
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text - remove ** markers
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        else:
            # Regular text
            if part:  # Only add non-empty parts
                run = paragraph.add_run(part)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

async def generate_contract(data: dict):
    """Generate a contract document using GPT-4o and save as .docx with logo"""
    try:
        deliverables_list = ", ".join(data.get("deliverables", []))
        terms_list = ", ".join(data.get("terms_conditions", []))
        
        document_content = ""
        async for chunk in contract_chain.astream({
            "contract_type": data.get("contract_type"),
            "party1_name": data.get("party1_name"),
            "party1_address": data.get("party1_address"),
            "party2_name": data.get("party2_name"),
            "party2_address": data.get("party2_address"),
            "service_description": data.get("service_description"),
            "contract_value": data.get("contract_value"),
            "payment_terms": data.get("payment_terms"),
            "duration": data.get("duration"),
            "deliverables": deliverables_list,
            "terms_conditions": terms_list,
            "effective_date": data.get("effective_date")
        }):
            document_content += chunk
        
        # Create professional DOCX
        contract_title = f"{data.get('contract_type', 'Service')} Contract"
        doc = create_professional_docx(document_content, contract_title)

        # Add logo to header using logo URL from request body
        logo_url = data.get("logo_url")
        if logo_url:
            print(f"Logo URL found in request: {logo_url}")
            add_logo_to_header(doc, logo_url)
        else:
            print("No logo URL provided in request data")

        # Save DOCX to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        # Generate unique filename
        unique_id = str(uuid.uuid4())[:8]
        filename = f"{data.get('contract_type', 'Service')}_Contract_{data.get('party1_name', '').replace(' ', '_')}_{data.get('party2_name', '').replace(' ', '_')}_{unique_id}.docx"

        # Upload DOCX to GCS
        client = storage.Client()
        bucket = client.bucket("deck123")
        docx_blob = bucket.blob(filename)
        docx_blob.upload_from_string(docx_bytes.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        docx_blob.make_public()
        document_url = docx_blob.public_url
        
        return {
            "document_content": document_content.strip(),
            "document_url": document_url,
            "document_type": f"{data.get('contract_type', 'Service')} Contract",
            "generated_for": f"{data.get('party1_name')} & {data.get('party2_name')}",
            "creation_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "word_count": len(document_content.split()),
            "format": "DOCX with logo"
        }
    except Exception as e:
        print(f"Error in contract generation: {e}")
        raise e
